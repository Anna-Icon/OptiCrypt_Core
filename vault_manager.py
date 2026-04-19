import os
from docx import Document
from pypdf import PdfReader
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


class VaultManager:
    def __init__(self, process_type):
        self.ptype = process_type  # enc / dec

    # ================= READ FILE =================
    def read_file(self, file_path):
        ext = os.path.splitext(file_path)[1].lower()

        # ===== TXT =====
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()

        # ===== WORD =====
        elif ext == ".docx":
            doc = Document(file_path)
            full_text = [paragraph.text for paragraph in doc.paragraphs]
            return "\n".join(full_text)

        # ===== PDF =====
        elif ext == ".pdf":
            reader = PdfReader(file_path)
            text = ""
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
            return text.strip()

        else:
            raise Exception("❌ Unsupported file type")

    # ================= SAVE FILE =================
    def save_file(self, original_path, content):
        name, ext = os.path.splitext(original_path)

        if name.endswith("_enc") or name.endswith("_dec"):
            name = name[:-4]

        new_path = f"{name}_{self.ptype}{ext}"

        # ================= TXT =================
        if ext == ".txt":
            with open(new_path, "w", encoding="utf-8") as f:
                f.write(content)

        # ================= WORD =================
        elif ext == ".docx":
            doc = Document()
            for line in content.split("\n"):
                doc.add_paragraph(line)
            doc.save(new_path)

        # ================= PDF (PRO VERSION) =================
        elif ext == ".pdf":
            self._save_pdf(content, new_path)

        else:
            raise Exception("❌ Unsupported file type")

        return new_path

    # ================= PDF WRITER =================
    def _save_pdf(self, content, path):
        c = canvas.Canvas(path, pagesize=letter)
        width, height = letter

        lines = content.split("\n")
        y = height - 40

        for line in lines:
            if y < 40:  # new page
                c.showPage()
                y = height - 40

            c.drawString(40, y, line[:100])  # حماية من overflow
            y -= 15

        c.save()

        # VaultManager
    def save_record(self, record):
            import json, os
            path = "central_vault.json"
            data = []

            # 1. التأكد من وجود الملف وقراءته
            if os.path.exists(path):
                with open(path, "r") as f:
                    try:
                        data = json.load(f)
                    except:
                        data = []

            # 2. إضافة السجل الجديد (ID, Key, Agent)
            data.append(record)

            # 3. الحفظ الفعلي على القرص
            with open(path, "w") as f:
                json.dump(data, f, indent=4)
            print(f"System: Record saved to {path}")
